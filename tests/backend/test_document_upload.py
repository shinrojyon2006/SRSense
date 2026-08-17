"""
Backend Document Upload & Text Extraction Test Suite.

Verifies PDF, DOCX, and TXT upload, storage, text extraction,
file validation (size & extension), and project ownership security.
"""

import io
import time
import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from httpx import AsyncClient, ASGITransport
from app.main import app


@pytest.mark.asyncio
async def test_document_upload_and_extraction_suite():
    """Test full document lifecycle: upload TXT, PDF, DOCX, extraction, security, and deletion."""
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        # 1. Register User A & Create Project
        email_a = f"doc_owner_{int(time.time() * 1000)}@srsense.ai"
        reg_a = await ac.post(
            "/api/auth/register",
            json={
                "name": "Doc Owner",
                "email": email_a,
                "password": "Password123!",
                "password_confirmation": "Password123!",
                "role": "developer",
            },
        )
        assert reg_a.status_code == 201
        token_a = reg_a.json()["access_token"]
        headers_a = {"Authorization": f"Bearer {token_a}"}

        proj_a = await ac.post(
            "/api/projects",
            json={"title": "SRS Ingestion Project", "description": "Document testing", "status": "active"},
            headers=headers_a,
        )
        assert proj_a.status_code == 201
        project_id = proj_a.json()["id"]

        # 2. Test Plain Text (.txt) Upload & Extraction
        txt_content = b"SRS Specification Document\n1. The SRSense System shall process requests within 200 milliseconds."
        txt_resp = await ac.post(
            f"/api/projects/{project_id}/documents/upload",
            files={"file": ("srs_spec.txt", txt_content, "text/plain")},
            headers=headers_a,
        )
        assert txt_resp.status_code == 201
        doc_txt_id = txt_resp.json()["id"]
        assert txt_resp.json()["status"] == "extracted"

        # Verify extracted text preview
        preview_res = await ac.get(
            f"/api/projects/{project_id}/documents/{doc_txt_id}/text",
            headers=headers_a,
        )
        assert preview_res.status_code == 200
        assert "process requests within 200 milliseconds" in preview_res.json()["raw_text"]

        # 3. Test PDF (.pdf) Upload & Extraction
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        pdf_buffer = io.BytesIO()
        writer.write(pdf_buffer)
        pdf_bytes = pdf_buffer.getvalue()

        pdf_resp = await ac.post(
            f"/api/projects/{project_id}/documents/upload",
            files={"file": ("architecture_doc.pdf", pdf_bytes, "application/pdf")},
            headers=headers_a,
        )
        assert pdf_resp.status_code == 201
        assert pdf_resp.json()["file_type"] == "pdf"

        # 4. Test DOCX (.docx) Upload & Extraction
        docx_doc = DocxDocument()
        docx_doc.add_heading("System Architecture Requirements", level=1)
        docx_doc.add_paragraph("The SRSense System shall store uploaded documents securely.")
        docx_buffer = io.BytesIO()
        docx_doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        docx_resp = await ac.post(
            f"/api/projects/{project_id}/documents/upload",
            files={
                "file": (
                    "spec.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            headers=headers_a,
        )
        assert docx_resp.status_code == 201
        doc_docx_id = docx_resp.json()["id"]

        # Verify DOCX extracted text
        docx_preview = await ac.get(
            f"/api/projects/{project_id}/documents/{doc_docx_id}/text",
            headers=headers_a,
        )
        assert docx_preview.status_code == 200
        assert "store uploaded documents securely" in docx_preview.json()["raw_text"]

        # 5. Test Invalid File Extension Rejection (.exe)
        invalid_resp = await ac.post(
            f"/api/projects/{project_id}/documents/upload",
            files={"file": ("malware.exe", b"binary content", "application/octet-stream")},
            headers=headers_a,
        )
        assert invalid_resp.status_code == 400
        assert "Unsupported file extension" in invalid_resp.json()["detail"]

        # 6. Test File Size Limit Exceeded Rejection (> 10 MB)
        oversized_content = b"X" * (10 * 1024 * 1024 + 1024)
        large_resp = await ac.post(
            f"/api/projects/{project_id}/documents/upload",
            files={"file": ("huge_spec.txt", oversized_content, "text/plain")},
            headers=headers_a,
        )
        assert large_resp.status_code == 400
        assert "exceeds maximum limit" in large_resp.json()["detail"]

        # 7. Test Project Ownership Security (User B blocked from User A's project)
        email_b = f"unauthorized_{int(time.time() * 1000)}@srsense.ai"
        reg_b = await ac.post(
            "/api/auth/register",
            json={
                "name": "User B",
                "email": email_b,
                "password": "Password123!",
                "password_confirmation": "Password123!",
                "role": "developer",
            },
        )
        assert reg_b.status_code == 201
        token_b = reg_b.json()["access_token"]
        headers_b = {"Authorization": f"Bearer {token_b}"}

        unauth_get = await ac.get(
            f"/api/projects/{project_id}/documents",
            headers=headers_b,
        )
        assert unauth_get.status_code == 404

        # 8. Test Document Deletion
        del_resp = await ac.delete(
            f"/api/projects/{project_id}/documents/{doc_txt_id}",
            headers=headers_a,
        )
        assert del_resp.status_code == 204
