"""
DOCX Text Extractor implementation using python-docx.
"""

from pathlib import Path
import docx
from app.core.extractors.base import BaseTextExtractor, ExtractedDocumentText, TextSegment


class DOCXTextExtractor(BaseTextExtractor):
    """Extracts text paragraph-by-paragraph and heading-by-heading from DOCX files."""

    def extract_text(self, file_path: Path) -> ExtractedDocumentText:
        doc = docx.Document(str(file_path))
        segments = []
        full_text_list = []
        current_section = "Document General"

        for idx, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if not text:
                continue

            # Check if paragraph is a heading
            if para.style and para.style.name.startswith("Heading"):
                current_section = text

            segments.append(
                TextSegment(location_label=f"Paragraph {idx} ({current_section})", text=text)
            )
            full_text_list.append(text)

        raw_text = "\n".join(full_text_list)
        words = raw_text.split()

        return ExtractedDocumentText(
            raw_text=raw_text,
            total_segments=len(segments),
            word_count=len(words),
            segments=segments,
            metadata={"paragraph_count": len(doc.paragraphs)},
        )
